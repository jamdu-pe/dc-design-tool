"""산출물 테스트 (Phase 4): xlsx·docx에 규격검증 결과가 실리는지 확인."""
import pytest
from docx import Document
from openpyxl import load_workbook

from dc_design_tool.engine.models import Spec
from dc_design_tool.engine.sizing import size
from dc_design_tool.reports.bom_xlsx import write_bom
from dc_design_tool.reports.design_basis_docx import write_design_basis


@pytest.fixture(scope="module")
def result():
    """Tier IV를 N+1로 설계 — 의도적 위반이 발생하는 케이스."""
    return size(Spec(project="report-test", rack_id="nvidia_gb200_nvl72",
                     it_power_mw=5.0, tier="IV", electrical_redundancy="N+1"))


def test_xlsx_has_compliance_sheet_listing_findings(result, tmp_path):
    wb = load_workbook(write_bom(result, str(tmp_path)))
    assert "규격검증" in wb.sheetnames
    rows = list(wb["규격검증"].values)
    assert rows[0] == ("심각도", "코드", "도메인", "판정", "설계값", "요구값", "근거")
    codes = {r[1] for r in rows[1:]}
    assert "TIER_ELECTRICAL" in codes


def test_xlsx_summary_reports_violation_count(result, tmp_path):
    wb = load_workbook(write_bom(result, str(tmp_path)))
    summary = {r[0]: r[1] for r in wb["요약"].values if r and r[0]}
    assert summary["규격 위반 건수"] == len(result.compliance.violations)
    assert summary["규격 위반 건수"] >= 1


def test_xlsx_bom_includes_network_cables(result, tmp_path):
    wb = load_workbook(write_bom(result, str(tmp_path)))
    items = {r[1] for r in wb["BOM"].values}
    assert "광케이블" in items


def test_docx_has_compliance_section_with_violation(result, tmp_path):
    doc = Document(write_design_basis(result, str(tmp_path)))
    text = "\n".join(p.text for p in doc.paragraphs)
    tables = "\n".join(c.text for t in doc.tables for row in t.rows for c in row.cells)
    assert "규격검증" in text
    assert "TIER_ELECTRICAL" in tables


def test_docx_still_contains_licensed_engineer_disclaimer(result, tmp_path):
    doc = Document(write_design_basis(result, str(tmp_path)))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "면허" in text


def test_docx_space_section_uses_new_area_breakdown(result, tmp_path):
    doc = Document(write_design_basis(result, str(tmp_path)))
    tables = "\n".join(c.text for t in doc.tables for row in t.rows for c in row.cells)
    assert "전기실" in tables and "기계실" in tables and "총 건축면적" in tables
