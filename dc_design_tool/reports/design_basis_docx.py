"""설계기준서(Design Basis) Word 산출. python-docx 기반."""
from __future__ import annotations
import pathlib
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from ..engine.models import SizingResult


def _h(doc, text, level=1):
    doc.add_heading(text, level=level)


def _kv_table(doc, rows):
    t = doc.add_table(rows=0, cols=2)
    t.style = "Light Grid Accent 1"
    for k, v in rows:
        c = t.add_row().cells
        c[0].text, c[1].text = str(k), str(v)
    return t


SEVERITY_LABEL = {"violation": "위반", "warning": "경고", "info": "정보"}
SEVERITY_ORDER = {"violation": 0, "warning": 1, "info": 2}


def _compliance_section(doc, result: SizingResult) -> None:
    """규격검증 결과를 심각도 순 표로 삽입."""
    report = result.compliance
    if report is None:
        doc.add_paragraph("규격검증이 수행되지 않았다.")
        return

    s = report.summary()
    doc.add_paragraph(
        f"Uptime Tier {report.tier} 및 ASHRAE 환경등급을 기준으로 검증한 결과, "
        f"위반 {s['violation']}건, 경고 {s['warning']}건, 정보 {s['info']}건이 확인되었다."
        + ("" if report.ok else " 위반 항목은 설계 변경 또는 등급 재설정이 필요하다."))

    t = doc.add_table(rows=1, cols=4)
    t.style = "Light Grid Accent 1"
    for cell, head in zip(t.rows[0].cells, ("심각도", "도메인", "판정", "근거")):
        cell.text = head
    for f in sorted(report.findings, key=lambda x: SEVERITY_ORDER[x.severity]):
        c = t.add_row().cells
        c[0].text = f"{SEVERITY_LABEL[f.severity]} [{f.code}]"
        c[1].text = f.domain
        c[2].text = f"{f.message} (설계 {f.actual} / 요구 {f.required})"
        c[3].text = f.rule


def write_design_basis(result: SizingResult, out_dir: str) -> str:
    out = pathlib.Path(out_dir)
    out.mkdir(parents=True, exist_ok=True)
    doc = Document()

    # 표지
    title = doc.add_heading("데이터센터 M&E 설계기준서 (Design Basis)", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(f"프로젝트: {result.project}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2 = doc.add_paragraph("개념설계 / 타당성 검토 수준")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # 1. 개요
    _h(doc, "1. 개요")
    doc.add_paragraph(
        f"본 문서는 {result.rack_id} 랙 {result.rack_count}대, 총 IT부하 "
        f"{result.it_power_kw} kW 규모의 데이터센터에 대한 기계·전기·통신(M&E) "
        "인프라 개념설계 기준을 정리한 것이다. 모든 수치는 결정론적 계산엔진으로 "
        "산출되었으며, 부품은 표준 인터페이스를 갖는 카탈로그 블록을 조합해 구성하였다.")

    # 2. 부하 집계
    _h(doc, "2. 부하 집계")
    _kv_table(doc, [
        ("랙 모델", result.rack_id),
        ("랙 수량", f"{result.rack_count} 대"),
        ("총 IT부하", f"{result.it_power_kw} kW"),
        ("설비 총부하(IT+냉각+하우스)", f"{result.electrical['facility_kw']} kW"),
        ("PUE(추정)", result.electrical["pue_estimate"]),
    ])

    # 3. 전기 설계 기준
    _h(doc, "3. 전기(Electrical) 설계 기준")
    e = result.electrical
    doc.add_paragraph(
        f"이중화 등급은 {e['redundancy']}를 적용한다. 무정전전원(UPS)은 배터리 "
        f"자립시간 {e['battery_autonomy_min']}분을 기준으로 발전기 기동·절체까지의 "
        "전원을 보장하며, 발전기는 비선형·스텝부하를 고려한 여유율을 반영한다. "
        f"수전은 {e['primary_kv']} kV 특고압을 전제로 한다.")
    _h(doc, "3.1 전원 계통 용량", level=2)
    _kv_table(doc, [
        ("UPS 필요용량", f"{e['ups_need_kva']} kVA"),
        ("UPS 구성", f"{e['ups_unit_kva']} kVA x {e['ups_qty']} 대 ({e['redundancy']})"),
        ("배터리 자립시간", f"{e['battery_autonomy_min']} 분"),
        ("배터리 필요에너지", f"{e['battery_energy_kwh']} kWh"),
        ("배터리 구성", f"{e['battery_qty']} 뱅크/캐비닛"),
        ("발전기 필요용량", f"{e['generator_need_kw']} kW"),
        ("발전기 구성", f"{e['generator_qty']} 대"),
        ("변압기 필요용량", f"{e['transformer_need_kva']} kVA"),
        ("변압기 구성", f"{e['transformer_qty']} 대"),
        ("수전 전압/전류", f"{e['primary_kv']} kV / {e['mv_current_a']} A"),
    ])
    _h(doc, "3.2 랙 급전 및 배전", level=2)
    _kv_table(doc, [
        ("랙 부하 전류", f"{e['rack_current_a']} A"),
        ("버스웨이 정격", f"{e['busway_rating_a']} A x {e['busway_qty']} 조"),
        ("랙 PDU 수량", f"{e['pdu_qty']} 대"),
        ("수용률", e["demand_factor"]),
    ])
    _h(doc, "3.3 고조파 대책", level=2)
    doc.add_paragraph(
        f"IT/UPS 등 비선형 부하의 전류 THD를 {int(e['thd_i_assumed']*100)}%로 "
        f"가정하고, 변압기는 고조파 여유율 x{e['harmonic_transformer_factor']}를 "
        "반영해 과열을 방지한다. 상세설계 시 K-factor 변압기 또는 능동필터 적용을 "
        "검토한다.")

    # 4. 기계 설계 기준
    _h(doc, "4. 기계(Mechanical, 냉각) 설계 기준")
    c = result.cooling
    doc.add_paragraph(
        f"IT 발열 {c['it_heat_kw']} kW 중 약 {int((c['liquid_kw']/c['it_heat_kw'])*100)}%를 "
        "직접칩냉각(D2C) 액냉이 흡수하고 잔열은 공냉으로 처리한다. "
        f"이중화 등급은 {c['redundancy']}를 적용한다.")
    _kv_table(doc, [
        ("총 발열", f"{c['it_heat_kw']} kW"),
        ("액냉 열량", f"{c['liquid_kw']} kW"),
        ("공냉 열량", f"{c['air_kw']} kW"),
        ("냉각수 유량", f"{c['coolant_flow_lpm']} L/min"),
        ("총 냉동톤", f"{c['total_rt']} RT"),
        ("CDU 구성", f"{c['cdu_qty']} 대 ({c['redundancy']})"),
        ("칠러 구성", f"{c['chiller_qty']} 대"),
    ])

    # 5. 통신 설계 기준
    _h(doc, "5. 통신(ICT) 설계 기준")
    n = result.network
    doc.add_paragraph(
        f"스케일아웃 패브릭은 리프-스파인 구조를 적용하며, 총 {n['scaleout_ports']} 포트"
        f"({n['port_speed_gbps']}G)를 수용한다. 구조화 배선은 TIA-942 기준을 준용한다.")
    _kv_table(doc, [
        ("스케일아웃 포트", f"{n['scaleout_ports']} p"),
        ("Leaf 스위치", f"{n['leaf_qty']} 대"),
        ("Spine 스위치", f"{n['spine_qty']} 대"),
        ("트랜시버", f"{n['transceiver_qty']} 개"),
    ])

    # 6. 공간/구조
    _h(doc, "6. 공간 및 구조")
    s = result.space
    doc.add_paragraph(
        f"랙은 열(row)당 {s['racks_per_row']}대 기준 {s['rack_rows']}개 열로 배치하며, "
        f"유효 층고는 액냉 배관·전기 트레이를 포함해 {s['clear_height_mm']} mm를 확보한다.")
    _kv_table(doc, [
        ("랙 점유면적", f"{s['rack_footprint_m2']} m2"),
        ("화이트스페이스(통로포함)", f"{s['white_space_m2']} m2"),
        ("전기실", f"{s['electrical_room_m2']} m2 "
                 f"(실내장비 {s['electrical_equipment_m2']} m2 x 이격 "
                 f"{s['equipment_clearance_factor']})"),
        ("기계실", f"{s['mechanical_room_m2']} m2 "
                 f"(실내장비 {s['mechanical_equipment_m2']} m2 x 이격 "
                 f"{s['equipment_clearance_factor']})"),
        ("지원공간(운영·보안·창고)", f"{s['support_area_m2']} m2"),
        ("총 건축면적", f"{s['total_building_m2']} m2"),
        ("랙 배치", f"{s['rack_rows']} 열 x {s['racks_per_row']} 대"),
        ("바닥 하중", f"{s['floor_load_kg_per_m2']} kg/m2 "
                    f"(허용 {s['floor_load_limit_kg_per_m2']} kg/m2)"),
    ])

    # 7. 규격검증
    _h(doc, "7. 규격검증 (Compliance)")
    _compliance_section(doc, result)

    # 8. 가정 및 리스크
    _h(doc, "8. 설계 가정 및 리스크")
    for a in result.assumptions:
        doc.add_paragraph(a, style="List Bullet")
    for w in result.warnings:
        doc.add_paragraph(f"[경고] {w}", style="List Bullet")

    # 고지문
    _h(doc, "9. 고지")
    note = doc.add_paragraph(
        "본 설계기준서는 개념설계/타당성 검토 수준의 산출물이며, 실시설계 및 "
        "인허가는 정보통신·전기·기계 분야 면허 기술자의 검토를 거쳐야 한다. "
        "미출시 장비(projected) 사양은 추정치이므로 벤더 확정값으로 갱신해야 한다.")
    note.runs[0].font.color.rgb = RGBColor(0x99, 0x00, 0x00)
    note.runs[0].font.size = Pt(9)

    path = out / "설계기준서.docx"
    doc.save(path)
    return str(path)
